from django.shortcuts import render
from django.http import JsonResponse
from django.conf import settings
from django.utils.timezone import now
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import os

from CarMechanic.buses.models import Modifications, Bus


def generate_word_file(request):
    if request.method == 'POST':
        brand = request.POST.get('brand', '').strip()
        number = request.POST.get('number', '').strip()
        kilometers = request.POST.get('kilometers', '').strip()
        dynamic_inputs = []
        sum_inputs = 0

        # Process dynamic inputs
        for key in request.POST:
            if key.startswith('text_input_'):
                index = key.split('_')[-1]
                text_value = request.POST.get(key, '').strip()
                number_value = request.POST.get(f'text_input_space_{index}', '0').strip()
                if not text_value.isdigit():
                    dynamic_inputs.append((text_value, int(number_value)))
                    sum_inputs += int(number_value)

        # Save data to the database
        try:
            bus, created = Bus.objects.get_or_create(
                model=brand,
                number=number,
                defaults={'km': int(kilometers) if kilometers.isdigit() else None},
            )

            for repair, price in dynamic_inputs:
                Modifications.objects.create(
                    bus=bus,
                    date=now().date(),
                    repair=repair,
                    price=price,
                )
        except Exception as e:
            return JsonResponse({'message': f'Error saving to database: {e}'}, status=500)

        # Generate the Word file
        if ' ' in number:
            filename_parts = number.split(' ')
            docx_filename = f"{filename_parts[-1]}.docx"
        else:
            docx_filename = f"{number}.docx"

        docx_filepath = os.path.join(settings.MEDIA_ROOT, docx_filename)

        if os.path.exists(docx_filepath):
            doc = Document(docx_filepath)
        else:
            doc = Document()

        brand_paragraph = doc.add_paragraph()
        brand_paragraph.paragraph_format.space_after = Pt(3)
        brand_paragraph.paragraph_format.space_before = Pt(3)
        brand_run = brand_paragraph.add_run(f'{brand} {number}')
        brand_run.font.size = Pt(14)
        brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_run.bold = True

        kilometers_paragraph = doc.add_paragraph()
        kilometers_paragraph.paragraph_format.space_before = Pt(3)
        kilometers_paragraph.paragraph_format.space_after = Pt(3)
        kilometers_run = kilometers_paragraph.add_run(f'{kilometers} км.')
        kilometers_run.font.size = Pt(14)
        kilometers_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if dynamic_inputs:
            p = doc.add_heading()
            heading_run = p.add_run('Извършен ремонт:')
            heading_run.font.size = Pt(14)
            heading_run.font.color.rgb = RGBColor(0, 0, 0)
            for i, (text_value, number_value) in enumerate(dynamic_inputs, start=1):
                table = doc.add_table(rows=1, cols=2)
                cell1 = table.cell(0, 0)
                cell1.width = Inches(6)
                cell1_paragraph = cell1.paragraphs[0]
                cell1_paragraph.paragraph_format.space_after = Pt(0)
                cell1_paragraph.paragraph_format.space_before = Pt(0)
                run1 = cell1_paragraph.add_run(f'{i}: {text_value}')
                run1.font.size = Pt(14)

                cell2 = table.cell(0, 1)
                cell2.width = Inches(0.9)
                cell2_paragraph = cell2.paragraphs[0]
                cell2_paragraph.paragraph_format.space_after = Pt(0)
                cell2_paragraph.paragraph_format.space_before = Pt(0)
                cell2_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run2 = cell2_paragraph.add_run(f'{number_value} лв.')
                run2.font.size = Pt(14)

        sum_paragraph = doc.add_paragraph()
        sum_paragraph.paragraph_format.space_after = Pt(3)
        sum_paragraph.paragraph_format.space_before = Pt(3)
        sum_run = sum_paragraph.add_run(f'Всичко: {sum_inputs} лв.')
        sum_run.bold = True
        sum_run.font.size = Pt(14)
        sum_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(docx_filepath)

        return render(request, 'wordfile/repairs.html', {
            'success_message': 'Генериран е WORD файл и данните са записани в базата!',
        })

    return render(request, 'wordfile/repairs.html')
