from datetime import datetime
from decimal import Decimal

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.conf import settings
from django.urls import reverse_lazy
from django.utils import timezone
from django.utils.timezone import now
from django.views.generic import ListView, DetailView
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import os

from CarMechanic.buses.models import Bus, Modifications
from CarMechanic.repairs.models import RepairSession


def generate_word_file(request):
    if request.method == 'POST':
        brand = request.POST.get('brand', '').strip()
        number = request.POST.get('number', '').strip()
        kilometers = request.POST.get('kilometers', '').strip()
        date_str = request.POST.get('date', '').strip()
        dynamic_inputs = []
        sum_inputs = Decimal(0)

        # Process dynamic inputs
        for key in request.POST:
            if key.startswith('text_input_'):
                index = key.split('_')[-1]
                text_value = request.POST.get(key, '').strip()
                number_value = Decimal(request.POST.get(f'text_input_space_{index}', '0').strip() or '0')
                if not text_value.isdigit():
                    dynamic_inputs.append((text_value, number_value))
                    sum_inputs += number_value

        # Save data to the database
        try:
            bus, created = Bus.objects.get_or_create(
                model=brand,
                number=number,
                defaults={'km': int(kilometers) if kilometers.isdigit() else None},
            )
            if date_str:
                try:
                    # Try to parse the date in DD.MM.YYYY format
                    date = datetime.strptime(date_str, '%Y-%m-%d')  # Django needs YYYY-MM-DD
                except ValueError:
                    # Handle the case where the date format is incorrect
                    return JsonResponse({'message': 'Invalid date format. Please use DD.MM.YYYY.'}, status=400)
            else:
                date = timezone.now()
            # Create a repair session
            repair_session = RepairSession.objects.create(bus=bus, total_cost=sum_inputs, date=date, km=kilometers)

            if kilometers.isdigit():
                kilometers = int(kilometers)
                if bus.km is None:
                    bus.km = kilometers
                    bus.save()
                elif bus.km < kilometers:
                    bus.km = kilometers
                    bus.save()

                    # Add modifications to the session
            for repair, price in dynamic_inputs:
                Modifications.objects.create(session=repair_session, repair=repair, price=price)
        except Exception as e:
            return JsonResponse({'message': f'Error saving to database: {e}'}, status=500)

        # Generate the Word file
        if ' ' in number:
            filename_parts = number.split(' ')
            docx_filename = f"{filename_parts[-1]}.docx"
        else:
            docx_filename = f"{number}.docx"

        docx_filepath = os.path.join(settings.MEDIA_ROOT, docx_filename)

        if os.path.exists(docx_filepath):
            doc = Document(docx_filepath)
        else:
            doc = Document()

        brand_paragraph = doc.add_paragraph()
        brand_paragraph.paragraph_format.space_after = Pt(3)
        brand_paragraph.paragraph_format.space_before = Pt(3)
        brand_run = brand_paragraph.add_run(f'{brand} {number}')
        brand_run.font.size = Pt(14)
        brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_run.bold = True

        kilometers_paragraph = doc.add_paragraph()
        kilometers_paragraph.paragraph_format.space_before = Pt(3)
        kilometers_paragraph.paragraph_format.space_after = Pt(3)
        kilometers_run = kilometers_paragraph.add_run(f'{kilometers} км.')
        kilometers_run.font.size = Pt(14)
        kilometers_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if dynamic_inputs:
            p = doc.add_heading()
            heading_run = p.add_run('Извършен ремонт:')
            heading_run.font.size = Pt(14)
            heading_run.font.color.rgb = RGBColor(0, 0, 0)
            for i, (text_value, number_value) in enumerate(dynamic_inputs, start=1):
                table = doc.add_table(rows=1, cols=2)
                cell1 = table.cell(0, 0)
                cell1.width = Inches(6)
                cell1_paragraph = cell1.paragraphs[0]
                cell1_paragraph.paragraph_format.space_after = Pt(0)
                cell1_paragraph.paragraph_format.space_before = Pt(0)
                run1 = cell1_paragraph.add_run(f'{i}: {text_value}')
                run1.font.size = Pt(14)

                cell2 = table.cell(0, 1)
                cell2.width = Inches(0.9)
                cell2_paragraph = cell2.paragraphs[0]
                cell2_paragraph.paragraph_format.space_after = Pt(0)
                cell2_paragraph.paragraph_format.space_before = Pt(0)
                cell2_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run2 = cell2_paragraph.add_run(f'{number_value} лв.')
                run2.font.size = Pt(14)

        sum_paragraph = doc.add_paragraph()
        sum_paragraph.paragraph_format.space_after = Pt(3)
        sum_paragraph.paragraph_format.space_before = Pt(3)
        sum_run = sum_paragraph.add_run(f'Всичко: {sum_inputs} лв.')
        sum_run.bold = True
        sum_run.font.size = Pt(14)
        sum_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(docx_filepath)

        return render(request, 'repairs/repairs.html', {
            'success_message': 'Генериран е WORD файл и данните са записани в базата!',
        })

    return render(request, 'repairs/repairs.html')


class BusRepairListView(ListView):
    model = RepairSession
    template_name = 'repairs/repair-list.html'
    context_object_name = 'repairs'  # This will be used in the template
    paginate_by = 5

    def get_queryset(self):
        # Get the bus by the ID passed in the URL
        bus = get_object_or_404(Bus, pk=self.kwargs['bus_id'])

        # Get all repair sessions for the bus
        return RepairSession.objects.filter(bus=bus)

    def get_context_data(self, **kwargs):
        # Get the default context data
        context = super().get_context_data(**kwargs)

        # Get the bus object again for the context
        bus = get_object_or_404(Bus, pk=self.kwargs['bus_id'])

        # Add the bus object to the context
        context['bus'] = bus

        return context


class RepairDetailView(DetailView):
    model = RepairSession
    template_name = 'repairs/repair-detail.html'
    context_object_name = 'repair_session'  # Make it clear that this is a RepairSession object

    def get_context_data(self, **kwargs):
        # Get the default context data
        context = super().get_context_data(**kwargs)

        # Get the current `RepairSession` object
        repair_session = self.get_object()

        # Add additional context
        context['repair_session'] = repair_session
        context['bus'] = repair_session.bus
        context['modifications'] = repair_session.modifications.all()

        # Add referer to the context
        referer = self.request.META.get('HTTP_REFERER', None)
        context['referer'] = referer if referer else reverse_lazy('bus_repairs', args=[repair_session.bus.id])

        return context